import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "submission.md"
TEMPLATE = ROOT / "docs" / "submission-guide" / "Copy of Digital Minds Research Sprint submission template.docx"
OUTPUT = ROOT / "report" / "Does_AI_wellbeing_survive_translation.docx"
SKILL_SCRIPTS = Path("/Users/ayesha/.codex/plugins/cache/openai-primary-runtime/documents/26.813.12317/skills/documents/scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT = "Old Standard TT"
BODY_SIZE = Pt(10)
CAPTION_SIZE = Pt(9)
CODE_SIZE = Pt(8)
CONTENT_WIDTH_DXA = 9360


def shade(element, fill):
    properties = element.get_or_add_pPr() if element.tag == qn("w:p") else element.get_or_add_tcPr()
    shd = properties.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        properties.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(8.5)):
    text = text.replace("**", "").replace("`", "")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = size
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563A8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(https?://[^\s)]+\)?|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline(paragraph, text, *, default_size=BODY_SIZE):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.font.name = FONT
            run.font.size = default_size
        token = match.group(0)
        if token.startswith("http"):
            suffix = ""
            while token and token[-1] in ".,;":
                suffix = token[-1] + suffix
                token = token[:-1]
            add_hyperlink(paragraph, token, token)
            if suffix:
                run = paragraph.add_run(suffix)
                run.font.name = FONT
                run.font.size = default_size
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = FONT
            run.font.size = default_size
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(default_size.pt - 0.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.name = FONT
            run.font.size = default_size
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.font.name = FONT
        run.font.size = default_size


def add_body_paragraph(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.05
    add_inline(paragraph, text)
    return paragraph


def add_code(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    shade(paragraph._p, "F2F2F2")
    for index, line in enumerate(lines):
        run = paragraph.add_run(("\n" if index else "") + line)
        run.font.name = "Consolas"
        run.font.size = CODE_SIZE
    return paragraph


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "B7B7B7")
        borders.append(border)
    table._tbl.tblPr.append(borders)
    first_weight = 1.55 if len(rows[0]) > 2 else 1.0
    weights = [first_weight] + [1.0] * (len(rows[0]) - 1)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            numeric = bool(re.fullmatch(r"[+\-−]?\d+(?:\.\d+)?%?", value.strip()))
            alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or numeric or column_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(
                table.cell(row_index, column_index),
                value,
                bold=row_index == 0,
                align=alignment,
            )
            if row_index == 0:
                shade(table.cell(row_index, column_index)._tc, "E9EEF5")
                table.rows[row_index]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_WIDTH_DXA)
    before = table._tbl.getprevious()
    if before is not None and before.tag == qn("w:p"):
        before_p = before
        p_pr = before_p.get_or_add_pPr()
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, image_path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        for doc_pr in drawing.iter(qn("wp:docPr")):
            doc_pr.set("descr", caption)
    paragraph.paragraph_format.keep_with_next = True
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(7)
    caption_p.paragraph_format.keep_with_next = True
    caption_run = caption_p.add_run(caption)
    caption_run.font.name = FONT
    caption_run.font.size = CAPTION_SIZE
    caption_run.italic = True


def parse_table(lines, index):
    table_lines = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = []
    for line_index, line in enumerate(table_lines):
        values = [cell.strip() for cell in line.strip("|").split("|")]
        if line_index == 1 and all(re.fullmatch(r":?-+:?", value) for value in values):
            continue
        rows.append(values)
    return rows, index


def source_body(source):
    return source.split("## 1. Introduction", 1)[1].join(["## 1. Introduction", ""])


def add_markdown_body(doc, source):
    body = "## 1. Introduction" + source.split("## 1. Introduction", 1)[1]
    lines = body.splitlines()
    index = 0
    paragraph_buffer = []

    def flush():
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush()
            index += 1
            continue
        if stripped.startswith("```"):
            flush()
            index += 1
            code = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            add_code(doc, code)
            continue
        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            flush()
            caption, relative = image_match.groups()
            add_figure(doc, (SOURCE.parent / relative).resolve(), caption)
            index += 1
            continue
        if stripped.startswith("|"):
            flush()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if stripped.startswith("# "):
            flush()
            heading = stripped[2:]
            if heading.startswith("Appendix A."):
                doc.add_page_break()
            paragraph = doc.add_paragraph(heading, style="Heading 2")
            index += 1
            continue
        if stripped.startswith("## "):
            flush()
            heading = stripped[3:]
            if heading in {"References"}:
                doc.add_page_break()
            paragraph = doc.add_paragraph(heading, style="Heading 2")
            index += 1
            continue
        if stripped.startswith("### "):
            flush()
            doc.add_paragraph(stripped[4:], style="Heading 3")
            index += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            flush()
            text = re.sub(r"^\d+\. ", "", stripped)
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, text)
            index += 1
            continue
        if stripped.startswith("- "):
            flush()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            flush()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(stripped[1:-1])
            run.italic = True
            run.font.name = FONT
            run.font.size = CAPTION_SIZE
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush()


def tune_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05
    for name in ("List Number", "List Bullet"):
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style.paragraph_format.space_after = Pt(3)
    for name, size in (("Heading 2", Pt(14)), ("Heading 3", Pt(11.5))):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = size
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(3)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(7)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(2)
    if "Caption" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = CAPTION_SIZE
        style.font.italic = True


def build():
    source = SOURCE.read_text()
    abstract = source.split("## Abstract\n\n", 1)[1].split("\n\n## 1.", 1)[0]
    doc = Document(TEMPLATE)
    tune_styles(doc)

    title_table = doc.tables[0]
    title_paragraph = title_table.cell(0, 0).paragraphs[1]
    title_paragraph.clear()
    title_run = title_paragraph.add_run("Does AI Wellbeing Survive Translation?")
    title_run.font.name = FONT
    title_run.font.size = Pt(20)
    title_run.bold = True

    title_cell = title_table.cell(1, 0)
    author_paragraph = title_cell.paragraphs[0]
    author_paragraph.clear()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_paragraph.paragraph_format.space_before = Pt(12)
    author_paragraph.paragraph_format.space_after = Pt(16)
    author_run = author_paragraph.add_run("Ayesha Imran  ·  Muhammad Aaliyan\n")
    author_run.font.name = FONT
    author_run.font.size = Pt(10.5)
    author_run.bold = True
    affiliation_run = author_paragraph.add_run("Independent  ·  Independent")
    affiliation_run.font.name = FONT
    affiliation_run.font.size = Pt(10)
    for nested_table in list(title_cell.tables):
        nested_table._element.getparent().remove(nested_table._element)
    apply_table_geometry(title_table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA)

    abstract_paragraph = title_cell.paragraphs[2]
    abstract_paragraph.clear()
    abstract_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_paragraph.paragraph_format.space_after = Pt(0)
    abstract_paragraph.paragraph_format.line_spacing = 1.05
    abstract_run = abstract_paragraph.add_run(abstract)
    abstract_run.font.name = FONT
    abstract_run.font.size = Pt(9.5)

    body = doc._element.body
    children = list(body)
    for child in children[3:-1]:
        body.remove(child)

    add_markdown_body(doc, source)
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    for element in doc._element.iter():
        width = element.get(qn("w:w"))
        if width and re.fullmatch(r"\d+\.0", width):
            element.set(qn("w:w"), str(int(float(width))))

    OUTPUT.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
